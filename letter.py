from docx import Document
from docx2pdf import convert

# Ambassador Program Letter Class
class LetterAP:
    # Constructor
    def __init__(self, firstName: str, lastName: str, credentials: str, volunteerHours: str, level: str, filePath: str):
        self.firstName = firstName
        self.lastName = lastName
        self.credentials = credentials
        self.filePath = filePath
        self.volunteerHours = volunteerHours
        self.level = level
    # Generate letter
    def generateLetter(self):
        # return -1 if the volunteer hours are not a number
        if not self.volunteerHours.isdigit():
            return -1

        # create word document object
        doc = Document()

        # add text to document piece by piece
        doc.add_paragraph('To whom it may concern,')


        firstParagraph = self.firstName + ' ' + self.lastName

        if self.level.lower() == 'regular':
            firstParagraph += ' is a '
        else:
            firstParagraph += ' is an '

        firstParagraph += self.level.lower() + ' Mètis ambassador for the Rupertsland Institute.'
        firstParagraph += ' To date, they have completed ' + str(self.volunteerHours) + ' hours of volunteer work.'

        doc.add_paragraph(firstParagraph)

        secondParagraph = 'In this role, ' + self.firstName + ' has accomplished the following:\n\n' + self.credentials

        doc.add_paragraph(secondParagraph)

        doc.add_paragraph('Thank you for your time and consideration,\n\n Person Who Signes Off.')

        # save document to specified file path
        doc.save(self.filePath + '/ambassador-letter-' + self.firstName.lower() + "-" + self.lastName.lower() + '.docx')

        # convert to pdf for emailing (dont want recipients to be able to edit the letter)
        convert(self.filePath + '/ambassador-letter-' + self.firstName.lower() + "-" + self.lastName.lower() + '.docx', self.filePath + '/ambassador-letter-' + self.firstName.lower() + "-" + self.lastName.lower() + '.pdf')

        return 0
